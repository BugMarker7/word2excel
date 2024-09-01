import os
import re
from docx import Document
from tqdm import tqdm
from openpyxl import Workbook
from utils.common import get_output_file


def get_template(template_file, omit_doc_name=False):
    '''
    Get the template of the docx file
    '''
    # word中的表格如果有使用合并单元格，python-docx读取会重复，使用text_map进行去重
    # 如果单元格内容不为空，就是要提取的位置，记录第几个表格，行，列，作为模板
    column_names = [] if omit_doc_name else ['']  # excel列名
    table_template, text_template, text_map = {}, {}, {}
    try:
        dfile = Document(template_file)
    except:
        raise ('打开模板文件失败!如果该文件在其他软件中打开，请关闭后再试!')

    # 获取表格模板，模板定义：在想被提取内容的单元格处填写 column_name
    tables = dfile.tables
    for tid, table in enumerate(tables):
        for rid, row in enumerate(table.rows):
            for cid, cell in enumerate(row.cells):
                text = cell.text.strip()
                if text != '' and text_map.get(text) == None:
                    key = f'{tid},{rid},{cid}'
                    table_template[key] = text
                    text_map[text] = key
                    column_names.append(text)

    # 获取文本模板，模板定义: 预提取位置填写 {{column_name}}
    paragraphs = dfile.paragraphs
    for pid, para in enumerate(paragraphs):
        text = para.text.strip()
        matched = re.findall(r'\s*{{[^}]+}}\s*', text)  # 提取每段所有{{column_name}}
        if len(matched) == 0:
            continue
        for m in matched:
            col_names = re.findall(r'{{([^}]+)}}', m)
            assert len(col_names) == 1, f'代码运行错误，请反馈:https://github.com/BugMarker7/word2excel'
            column_name = col_names[0].strip()
            if column_name == '':
                raise Exception(f'模板文件中存在空列名，请检查模板文件!')
            if text_map.get(column_name) != None:
                raise Exception(f'模板文件中存在重复的列名{column_name}，请检查模板文件!')
            text_map[column_name] = pid
            column_names.append(column_name)

            text = text.replace(m, '(.*)')
        text_template[pid] = (re.compile(text), len(matched))
    return table_template, text_template, column_names


def word2excel(template_file, docx_dir, output_dir='output', omit_doc_name=False):
    table_template, text_template, column_names = get_template(
        template_file, omit_doc_name)
    wb = Workbook()
    ws = wb.active
    ws.append(column_names)

    docx_files = [file for file in os.listdir(docx_dir) if not file.startswith('~$')]
    for docx_file in tqdm(docx_files):
        try:
            dfile = Document(os.path.join(docx_dir, docx_file))
        except:
            print(f'打开文件 {docx_file} 失败!如果该文件在其他软件中打开，请关闭后再试!')
            exit()

        data = [] if omit_doc_name else [docx_file]
        # 提取表格数据
        tables = dfile.tables
        for tid, table in enumerate(tables):
            for rid, row in enumerate(table.rows):
                for cid, cell in enumerate(row.cells):
                    key = f'{tid},{rid},{cid}'
                    if table_template.get(key) != None:
                        data.append(cell.text.strip())

        # 提取文本数据
        paragraphs = dfile.paragraphs
        for pid, para in enumerate(paragraphs):
            if text_template.get(pid) == None:
                continue
            text = para.text.strip()
            matched = text_template[pid][0].findall(text)
            if len(matched) != 1 and len(matched[0]) != text_template[pid][1]:
                print(f'{docx_file} 提取第 {pid + 1} 段文本数据出现错误，请检查该文档')
            values = [m.strip() for m in matched[0]]
            data.extend(values)
        ws.append(data)

    output_file = get_output_file(output_dir)
    wb.save(output_file)
    print(f'程序运行结束，Excel文件保存到 {os.path.realpath(output_file)}' )
